"""DOCX parsing strategy using python-docx."""

from __future__ import annotations

import asyncio
import io

import docx  # python-docx

from app.parsers.base_parser import BaseParser


class DocxParser(BaseParser):
    """Extracts text from Microsoft Word (.docx) resumes."""

    extension = ".docx"

    async def extract_text(self, file_bytes: bytes) -> str:
        """Extract text from DOCX bytes.

        python-docx is synchronous, so we run it in a thread to keep the async
        event loop responsive.
        """
        text = await asyncio.to_thread(self._extract, file_bytes)
        if not text.strip():
            raise ValueError("Could not extract any text from the DOCX file.")
        return self._normalize(text)

    @staticmethod
    def _extract(file_bytes: bytes) -> str:
        """Pull text from paragraphs AND tables. Synchronous/blocking."""
        document = docx.Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        # 1) Body paragraphs (the bulk of most resumes).
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        # 2) Tables — many resumes lay out skills/dates in tables, so we walk
        #    every row and cell to avoid losing that content.
        for table in document.tables:
            for row in table.rows:
                # Join non-empty cells in the row with a tab separator.
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append("\t".join(cells))

        return "\n".join(parts)
